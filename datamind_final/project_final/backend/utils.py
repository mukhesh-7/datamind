from PyPDF2 import PdfReader
import docx
import os
import pdfplumber
import pytesseract
from PIL import Image
import io

async def extract_text_from_file(file) -> str:
    try:
        if file.content_type == "application/pdf":
            # Try pdfplumber first
            try:
                file.file.seek(0)
                with pdfplumber.open(file.file) as pdf:
                    text = ""
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted
                    if text.strip():
                        return text
            except Exception as e:
                print(f"pdfplumber extraction error: {e}")
            # Fallback to PyPDF2
            try:
                file.file.seek(0)
                reader = PdfReader(file.file)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted
                if text.strip():
                    return text
            except Exception as e:
                print(f"PyPDF2 extraction error: {e}")
            # Fallback to OCR (for scanned PDFs)
            try:
                file.file.seek(0)
                with pdfplumber.open(file.file) as pdf:
                    text = ""
                    for page in pdf.pages:
                        img = page.to_image(resolution=300).original
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text:
                            text += ocr_text
                    if text.strip():
                        return text
            except Exception as e:
                print(f"OCR fallback error: {e}")
            return "Error: PDF contains no extractable text. (Tried text and OCR)"
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from DOCX
            try:
                file.file.seek(0)
                doc = docx.Document(file.file)
                text = "\n".join([para.text for para in doc.paragraphs])
                # Extract from tables too
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                text += '\n' + cell_text
                if not text.strip():
                    return "Error: DOCX contains no extractable text."
                return text
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                return f"Error: Unable to extract text from DOCX. {e}"
        elif file.content_type == "text/plain":
            # Extract text from plain text file
            file.file.seek(0)
            return file.file.read().decode("utf-8")
        elif file.content_type == "application/msword":
            return "Error: .doc (legacy Word) files are not supported. Please upload .docx files."
        else:
            raise ValueError("Unsupported file type.")
    except Exception as e:
        print(f"Error extracting text: {e}")
        return f"Error: Unable to extract text from the file. {e}"

# Utility to generate a short description from document content

def extract_description(content: str) -> str:
    """
    Generate a short description for a document based on its content.
    """
    if not content:
        return "No content available."
    # Take the first 30 words as a simple summary
    words = content.split()
    desc = " ".join(words[:30])
    if len(words) > 30:
        desc += "..."
    return desc
